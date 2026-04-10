"""
blob_loader.py
Loads documents from an Azure Blob Storage container.
Supports .txt, .md, .pdf, and .docx files.
"""
import io
import logging
from dataclasses import dataclass
from typing import Generator
 
from azure.storage.blob import BlobServiceClient
from azure.core.exceptions import AzureError
 
logger = logging.getLogger(__name__)
 
 
@dataclass
class BlobDocument:
    name: str          # blob name (path inside container)
    content: str       # extracted plain text
    metadata: dict     # source, size, last_modified, etc.
 
 
class BlobLoader:
    """
    Iterates over every blob in a container and yields BlobDocument objects
    with the raw text extracted from supported file types.
    """
 
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
 
    def __init__(self, connection_string: str, container_name: str):
        self.client = BlobServiceClient.from_connection_string(connection_string)
        self.container = self.client.get_container_client(container_name)
        self.container_name = container_name
 
    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------
 
    def load_all(self) -> Generator[BlobDocument, None, None]:
        """Yield one BlobDocument per supported blob in the container."""
        for blob_props in self.container.list_blobs():
            name: str = blob_props.name
            ext = self._extension(name)
            if ext not in self.SUPPORTED_EXTENSIONS:
                logger.debug("Skipping unsupported file: %s", name)
                continue
 
            try:
                raw_bytes = self._download(name)
                text = self._extract_text(raw_bytes, ext, name)
                if not text.strip():
                    logger.warning("Empty content for blob: %s", name)
                    continue
                yield BlobDocument(
                    name=name,
                    content=text,
                    metadata={
                        "source": f"azure-blob://{self.container_name}/{name}",
                        "size": blob_props.size,
                        "last_modified": str(blob_props.last_modified),
                        "content_type": blob_props.content_settings.content_type or "",
                    },
                )
            except Exception as exc:  # noqa: BLE001
                logger.error("Failed to process blob %s: %s", name, exc)
 
    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------
 
    def _download(self, blob_name: str) -> bytes:
        blob_client = self.container.get_blob_client(blob_name)
        return blob_client.download_blob().readall()
 
    @staticmethod
    def _extension(name: str) -> str:
        dot = name.rfind(".")
        return name[dot:].lower() if dot != -1 else ""
 
    def _extract_text(self, raw: bytes, ext: str, name: str) -> str:
        if ext in {".txt", ".md"}:
            return raw.decode("utf-8", errors="replace")
 
        if ext == ".pdf":
            return self._extract_pdf(raw)
 
        if ext == ".docx":
            return self._extract_docx(raw)
 
        return ""
 
    @staticmethod
    def _extract_pdf(raw: bytes) -> str:
        try:
            import pypdf  # lazy import — only needed if PDFs are present
 
            reader = pypdf.PdfReader(io.BytesIO(raw))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
        except ImportError:
            logger.error("pypdf not installed. Run: pip install pypdf")
            return ""
 
    @staticmethod
    def _extract_docx(raw: bytes) -> str:
        try:
            import docx  # python-docx
 
            doc = docx.Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return ""
 
 